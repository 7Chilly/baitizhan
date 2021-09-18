import openpyxl
import docx


def replaceInfo(doc, oldInfo, newInfo):
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = run.text.replace(oldInfo, newInfo)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace(oldInfo, newInfo)

wb = openpyxl.load_workbook("夜曲大学英语考试成绩.xlsx", data_only= True)
ws = wb["汇总"]
firstRow = ()
for rowIndex, row in enumerate(ws.rows):
    if rowIndex == 0:
        firstRow = row
    else:
        doc = docx.Document("成绩报告单模版.docx")
        for columnIndex, cell in enumerate(row):
            oldInfo = firstRow[columnIndex].value
            newInfo = str(cell.value)
            replaceInfo(doc, oldInfo, newInfo)
        name = row[0].value
        doc.save(f"学生成绩单/成绩报告单_{name}.docx")
