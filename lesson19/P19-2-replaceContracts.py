import docx
import openpyxl
import os

def replaceInfo(doc, oldInfo, newInfo):
    for para in doc.paragraphs:
        for run in para.runs:
                run.text = run.text.replace(oldInfo, newInfo)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace(oldInfo, newInfo)

wb = openpyxl.load_workbook("采购信息列表.xlsx", data_only= True)
ws = wb["6月采购"]
firstRow = ()
for rowIndex, row in enumerate(ws.rows):
    if rowIndex == 0:
        firstRow = row
    else:
        doc = docx.Document("合同模版.docx")
        for columnIndex, cell in enumerate(row):
            newInfo = str(cell.value)
            oldInfo = str(firstRow[columnIndex].value)
            replaceInfo(doc, oldInfo, newInfo)
        productName = row[3].value
        doc.save(f"合同/采购合同_{productName}.docx")
print("success")


