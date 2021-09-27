import openpyxl
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm,Pt
from docx.oxml.ns import qn

wb = openpyxl.load_workbook("招股书表格.xlsx", data_only= True)
doc = docx.Document("九、主要财务指标.docx")
for idx, ws in enumerate(wb.worksheets):
    for rowIndex, row in enumerate(ws.rows):
        table = doc.tables[idx]
        if rowIndex < len(table.rows):
            for columnIndex, cell in enumerate(row):
                if columnIndex < len(table.columns):
                    Info = cell.value
                    if type(Info) != int and type(Info) != float and Info != "/":
                        continue
                    if Info == "/":
                        Info = "不适用"
                        run = table.cell(rowIndex, columnIndex).paragraphs[0].add_run(str(Info))
                        run.font.name = "宋体"
                        run.font.size = Pt(10.5)
                        r = run._element
                        r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    if type(Info) == int or type(Info) == float:
                        Info = format(float(Info), '.2f')
                        Info = format(float(Info), ',')
                        a_len = len(Info.split('.')[1])
                        if a_len == 1:
                            Info += "0"
                        run = table.cell(rowIndex, columnIndex).paragraphs[0].add_run(str(Info))
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10.5)
                    if Info == None:
                        continue
                    table.cell(rowIndex, columnIndex).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.RIGHT

doc.save("九、主要财务指标_改.docx")



