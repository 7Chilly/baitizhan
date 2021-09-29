import openpyxl
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn


def getBold(table, rowIndex, columnIndex):
    firstCell = table.cell(rowIndex,0).paragraphs[0].runs[0].bold
    if columnIndex == 0:
        if firstCell:
            return True
    else:
        secondCell = table.cell(rowIndex,1).paragraphs[0].runs[0].bold
        if firstCell or secondCell:
            return True


def changeChinese(run):
    run.font.name = "宋体"
    run.font.size = Pt(10.5)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return run


def replaceValue(table, rowIndex, columnIndex, Info):
    originalValue = table.cell(rowIndex, columnIndex).text
    replaceValue = str(Info)
    table.cell(rowIndex, columnIndex).text = table.cell(rowIndex, columnIndex).text.replace(originalValue, replaceValue)


def twoDecPer(cell):
    cellValue = cell.value
    cellValue *= 100
    cellValue = f"{cellValue:.2f}"
    Info = str(cellValue) + "%"
    return Info

def zeroDecPer(cell):
    cellValue = cell.value
    cellValue *= 100
    cellValue = f"{cellValue:.0f}"
    Info = str(cellValue) + "%"
    return Info


wb = openpyxl.load_workbook("招股书表格.xlsx", data_only= True)
docName = "九、主要财务指标.docx"
doc = docx.Document(docName)
for idx, ws in enumerate(wb.worksheets):
    table = doc.tables[idx]
    for rowIndex, row in enumerate(ws.rows):
        if rowIndex < len(table.rows):
            for columnIndex, cell in enumerate(row):
                if columnIndex < len(table.columns):
                    Info = cell.value
                    if str(Info) == table.cell(rowIndex, columnIndex).paragraphs[0].text:
                        continue
                    if Info is None:
                        continue
                    if type(Info) != int and type(Info) != float:
                        if "%" in str(Info) or str(Info) == "/":
                            replaceValue(table, rowIndex, columnIndex, Info)
                            run = table.cell(rowIndex, columnIndex).paragraphs[0].runs[0]
                            run = changeChinese(run)
                            run.font.name = "Times New Roman"
                            if getBold(table, rowIndex):
                                table.cell(rowIndex, columnIndex).paragraphs[0].runs[0].font.bold = True
                            table.cell(rowIndex, columnIndex).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.RIGHT
                        else:
                            replaceValue(table, rowIndex, columnIndex, Info)
                            run = table.cell(rowIndex, columnIndex).paragraphs[0].runs[0]
                            run = changeChinese(run)
                            run.font.name = "Times New Roman"
                            if getBold(table, rowIndex):
                                table.cell(rowIndex, columnIndex).paragraphs[0].runs[0].font.bold = True
                            table.cell(rowIndex, columnIndex).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
                    if type(Info) == int or type(Info) == float:
                        formatCell = cell.number_format
                        if "%" in formatCell:
                            if formatCell == "0%":
                                Info = zeroDecPer(cell)
                            if formatCell == "0.00%":
                                Info = twoDecPer(cell)
                        else:
                            Info = format(float(Info), '.2f')
                            Info = format(float(Info), ',')
                            a_len = len(Info.split('.')[1])
                            if a_len == 1:
                                Info += "0"
                        if str(Info) == table.cell(rowIndex, columnIndex).paragraphs[0].text:
                            continue
                        replaceValue(table, rowIndex, columnIndex, Info)
                        run = table.cell(rowIndex, columnIndex).paragraphs[0].runs[0]
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10.5)
                        if getBold(table, rowIndex):
                            table.cell(rowIndex, columnIndex).paragraphs[0].runs[0].font.bold = True
                        table.cell(rowIndex, columnIndex).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.RIGHT
saveName = docName.split(".")[0] + "003.docx"
doc.save(saveName)



