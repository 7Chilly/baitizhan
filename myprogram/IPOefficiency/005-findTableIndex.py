import docx

doc = docx.Document("【有研硅】财务部分.docx")
checkInfo1 = input("请填写您想查找的表格内容（最好unique一点哦）：")
InfoList = []
for idx, table in enumerate(doc.tables):
    for row in table.rows:
        for cell in row.cells:
            cellValue = cell.text
            if checkInfo1 == cellValue:
                InfoList.append(idx)
if len(InfoList) == 1:
    idx = InfoList[0]
    print(f"这是Sheet{idx}")
else:
    checkInfo2 = input("信息可以再精准一点，再填一个试试吧：")
    newInfoList = []
    for idx in InfoList:
        table = doc.tables[idx]
        for row in table.rows:
            for cell in row.cells:
                cellValue = cell.text
                if checkInfo2 == cellValue:
                    newInfoList.append(idx)
    if len(newInfoList) == 1:
        idx = newInfoList[0]
        print(f"这是Sheet{idx}")
    else:
        if len(newInfoList) == 0:
            print("没有找到哦，请检查是否输入正确～")
        else:
            print("这是包含输入信息的table索引号，自己寻找吧！", newInfoList)
