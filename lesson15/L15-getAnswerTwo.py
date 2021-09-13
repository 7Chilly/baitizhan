import docx
import os

standardTwo = ["东临碣石", "行舟绿水前", "孤山寺北贾亭西", "断肠人在天涯", "故人具鸡黍", "一曲新词酒一杯", "何当共剪西窗烛", "误入藕花深处", "烟笼寒水月笼沙", "万籁此都寂", "初日照高林", "腾蛇乘雾"]

allItems = os.listdir()
allStudentData= []

for item in allItems:
    studentData = {}
    studentInfo = os.path.splitext(item)[0]
    studentData["Class"] = studentInfo.split("-")[0]
    studentData["Name"] = studentInfo.split("-")[1]
    doc = docx.Document(item)
    studentData["ID"] = doc.paragraphs[3].runs[1].text
    studentData["ScoreTwo"] = 0
    for idx,value in enumerate(standardTwo, 8):
        if doc.paragraphs[idx].runs[1].text == value:
            studentData["ScoreTwo"] += 5

    allStudentData.append(studentData)
print(allStudentData)
