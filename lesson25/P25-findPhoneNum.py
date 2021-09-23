# 不断打开文件夹，需要调用函数；
# word/excel/pdf读取方式不同，也需要调用函数；
# 输入：循环遍历所有word/excel/pdf文件，判断里面有没有电话号码；
# 输出：不重复的numList
# 杀人诛心啊

import os
import docx
import openpyxl
import pdfplumber
import re


def findNum(string):
    NumList = []
    pattern = re.compile(r"\d+")
    res = re.findall(pattern, string)
    for number in res:
        lenNumber = len(number)
        if lenNumber == 11:
            NumList.append(number)
    return NumList


def getList(searchPath):
    OriginalList = []
    for filepath, dirnames, filenames in os.walk(searchPath):
        for filename in filenames:
            itemPath = os.path.join(filepath, filename)
            extension = os.path.splitext(itemPath)[1].lower()

            if extension == ".pdf":
                pdf = pdfplumber.open(itemPath)
                for page in pdf.pages:
                    pageText = page.extract_text()
                    NumList = findNum(pageText)
                    if NumList is None:
                        continue
                    OriginalList += NumList

                for page in pdf.pages:
                    table = page.extract_tables()
                    for row in table:
                        for cell in row:
                            cellText = str(cell)
                            NumList = findNum(cellText)
                            if NumList is None:
                                continue
                            OriginalList += NumList

            if extension == ".docx":
                doc = docx.Document(itemPath)
                for para in doc.paragraphs:
                    for run in para.runs:
                        runText = run.text
                        NumList = findNum(runText)
                        if NumList is None:
                            continue
                        OriginalList += NumList
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            cellText = cell.text
                            NumList = findNum(cellText)
                            if NumList is None:
                                continue
                            OriginalList += NumList

            if extension == ".xlsx":
                wb = openpyxl.load_workbook(itemPath, data_only=True)
                for ws in wb.worksheets:
                    for row in ws.rows:
                        for cell in row:
                            cellValue = str(cell.value)
                            NumList = findNum(cellValue)
                            if NumList is None:
                                continue
                            OriginalList += NumList
    phoneList = list(set(OriginalList))
    print(phoneList)


getList("/Users/chilly/Desktop/python/yequ/再见啦/problem25")
