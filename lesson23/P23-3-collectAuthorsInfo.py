import os
import openpyxl
import docx

sourcePath = "sub"
allItems = os.listdir(sourcePath)
wb = openpyxl.load_workbook("作者投稿.xlsx")
ws = wb["6月"]
for item in allItems:
    itemPath = os.path.join(sourcePath, item)
    info = []
    doc = docx.Document(itemPath)
    article = doc.paragraphs[0].text
    info.append(article)
    authorInfo = doc.paragraphs[2].text
    information = authorInfo.split(" ")
    for i in information:
        info.append(i)
    ws.append(info)
wb.save("作者投稿.xlsx")
