import os
import docx
import openpyxl

sourcePath = "article"
allItems = os.listdir(sourcePath)
for item in allItems:
    itemPath = os.path.join(sourcePath, item)
    if os.path.splitext(item)[1] != ".docx":
        continue
    doc = docx.Document(itemPath)
    name = doc.paragraphs[2].runs[0].text
    email = doc.paragraphs[2].runs[3].text
    wb = openpyxl.load_workbook("作者投稿-20年.xlsx")
    for sheet in wb.worksheets:
        for row in sheet.rows:
            name1 = row[1].value
            email1 = row[3].value
            if name == name1 and email == email1:
                month = sheet.title
                doc.save(f"modify/2020年{month}{name}.docx")
                print(f"{name}修改完成")
