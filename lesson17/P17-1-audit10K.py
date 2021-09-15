import openpyxl
import docx

wb = openpyxl.load_workbook("年报.xlsx",data_only= True)
doc = docx.Document("年报.docx")
for idx, ws in enumerate(wb.worksheets):
    for rowIndex, row in enumerate(ws.rows):
        for columnIndex, cell in enumerate(row):
            Info = cell.value
            if type(Info) == int:
                if Info < 0:
                    Info = abs(Info)
                    Info = (f"{Info:,}")
                    Info = f"({Info})"
                else:
                    Info = (f"{Info:,}")
            if Info == None:
                continue
            doc.tables[idx].cell(rowIndex, columnIndex).paragraphs[0].runs[0].text = Info

doc.save("年报_改.docx")
print("success")
