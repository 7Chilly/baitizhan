import openpyxl
import docx

wb = openpyxl.load_workbook("夜曲大学英语考试成绩.xlsx", data_only= True)
ws = wb["汇总"]
for row in ws.rows:
    doc = docx.Document("成绩报告单模版.docx")

    name = row[0].value
    school = row[1].value
    college = row[2].value
    studentID = row[3].value
    certificateID = row[4].value
    listening = row[5].value
    reading = row[6].value
    writing = row[7].value
    total = row[8].value
    if name == "姓名":
        continue

    doc.paragraphs[3].runs[2].add_text(name)
    doc.paragraphs[4].runs[2].add_text(school)
    doc.paragraphs[5].runs[1].add_text(college)
    doc.paragraphs[6].runs[1].add_text(studentID)
    doc.paragraphs[10].runs[0].add_text(certificateID)

    doc.tables[0].cell(1, 0).text = str(total)
    doc.tables[0].cell(1, 1).text = str(listening)
    doc.tables[0].cell(1, 2).text = str(reading)
    doc.tables[0].cell(1, 3).text = str(writing)

    doc.save(f"学生成绩单/{name}的英语成绩单.docx")
print("success")
