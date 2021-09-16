import os
import docx
import openpyxl
from docx import shared

wb = openpyxl.load_workbook("志愿者名单.xlsx")
ws = wb["名单"]
for row in ws.rows:
    doc = docx.Document("志愿者工牌模版.docx")
    name = row[0].value
    team = row[1].value
    if name == "姓名":
        continue
    allPicPath = "photo"
    picPath = os.path.join(allPicPath, f"{name}.png")
    doc.paragraphs[7].runs[0].add_text(name)
    doc.paragraphs[8].runs[0].add_text(team)
    table = doc.tables[0]
    run = table.cell(0, 0).paragraphs[0].add_run("")
    run.add_picture(picPath, width=shared.Cm(5))
    doc.save(f"志愿者/工牌_{name}.docx")
print("success")
