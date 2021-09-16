import docx
import random

sourceDoc = docx.Document("题库.docx")
firstList = []
secondList = []
for i in range(0, 12):
    para = sourceDoc.paragraphs[i].text
    firstList.append(para)
for i in range(14, 22):
    para = sourceDoc.paragraphs[i].text
    secondList.append(para)

for i in range(0, 72):
    doc = docx.Document("试卷模版.docx")
    testQuestionList1 = random.sample(firstList, 6)
    testQuestionList2 = random.sample(secondList, 4)
    testQuestionList = testQuestionList1 + testQuestionList2
    for idx, question in enumerate(testQuestionList):
        doc.paragraphs[idx].add_run(question)
    doc.save(f"examination/试卷_{i+1}.docx")
print("success")

# remember to differentiate add_run from add_text
