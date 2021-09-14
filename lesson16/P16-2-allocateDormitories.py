import os
import docx
import shutil

os.mkdir("男")
os.mkdir("女")
yes = "需要安静的学习环境"
no = "不太需要安静的学习环境"
os.mkdir(f"男/{yes}")
os.mkdir(f"男/{no}")
os.mkdir(f"女/{yes}")
os.mkdir(f"女/{no}")
allItems = os.listdir()
for item in allItems:
    if os.path.isdir(item):
        continue
    doc = docx.Document(item)
    gender = doc.paragraphs[2].runs[1].text
    score = int(doc.tables[0].cell(2, 1).text)
    if score > 5:
        shutil.move(item, f"{gender}/{yes}")
    else:
        shutil.move(item, f"{gender}/{no}")
