import os
import openpyxl
import docx

sourcePath = "统计表"
allItems = os.listdir(sourcePath)
wb = openpyxl.load_workbook("我是健康小卫士打分表.xlsx")
sheet = wb["我是健康小卫士"]

for item in allItems:
    schoolInfo = {}
    filePath = os.path.join(sourcePath, item)
    doc = docx.Document(filePath)
    schoolInfo["school"] = doc.paragraphs[1].runs[1].text
    table = doc.tables[0]
    schoolInfo["studentName"] = table.cell(2, 2).text
    if schoolInfo["studentName"] == "":
        continue
    schoolInfo["classInfo"] = table.cell(2, 3).text
    schoolInfo["teacher"] = table.cell(2, 4).text
    sheet.append(list(schoolInfo.values()))

wb.save("我是健康小卫士打分表.xlsx")
